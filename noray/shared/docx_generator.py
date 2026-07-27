"""
NORAY — Executive Word Document (.docx) Generator

Creates professionally formatted Microsoft Word documents for CVs, SOPs,
Motivation Letters, and Research Proposals using python-docx.
"""

from __future__ import annotations

from pathlib import Path

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt, RGBColor

from noray.shared.models import CareerProfile


def set_cell_background(cell, fill_hex: str):
    """Set shading color for a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)


def add_heading_styled(doc, text: str, level: int = 1, color_rgb=(16, 185, 129)):
    """Add a styled heading with custom font size and color."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(*color_rgb)
    elif level == 2:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(31, 41, 55)
    return p


def generate_cv_docx(
    profile: CareerProfile,
    company: str,
    role: str,
    output_path: Path,
) -> Path:
    """
    Generate an executive ATS-optimized Microsoft Word (.docx) resume.
    """
    doc = docx.Document()

    # Set page margins (0.75 in)
    sections = doc.sections
    for s in sections:
        s.top_margin = Inches(0.75)
        s.bottom_margin = Inches(0.75)
        s.left_margin = Inches(0.75)
        s.right_margin = Inches(0.75)

    # 1. Header (Candidate Name & Contact Info)
    name = profile.identity.name or "Gharieb Mohamed"
    role_title = role or "Machine Learning & AI Engineer"

    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_header.paragraph_format.space_after = Pt(2)

    r_name = p_header.add_run(name)
    r_name.font.name = 'Calibri'
    r_name.font.size = Pt(24)
    r_name.font.bold = True
    r_name.font.color.rgb = RGBColor(16, 185, 129) # Emerald

    p_subtitle = doc.add_paragraph()
    p_subtitle.paragraph_format.space_after = Pt(6)
    r_sub = p_subtitle.add_run(f"{role_title} — Tailored for {company}")
    r_sub.font.name = 'Calibri'
    r_sub.font.size = Pt(12)
    r_sub.font.color.rgb = RGBColor(107, 114, 128) # Gray

    # Contact line
    contacts = []
    if profile.identity.email:
        contacts.append(profile.identity.email)
    if profile.identity.phone:
        contacts.append(profile.identity.phone)
    if profile.identity.location.country:
        loc = f"{profile.identity.location.city}, {profile.identity.location.country}".strip(", ")
        contacts.append(loc)
    if profile.identity.linkedin_url:
        contacts.append(profile.identity.linkedin_url)
    if profile.identity.github_url:
        contacts.append(profile.identity.github_url)

    if contacts:
        p_contact = doc.add_paragraph()
        p_contact.paragraph_format.space_after = Pt(12)
        r_c = p_contact.add_run(" | ".join(contacts))
        r_c.font.name = 'Calibri'
        r_c.font.size = Pt(9.5)
        r_c.font.color.rgb = RGBColor(75, 85, 99)

    # Horizontal Divider line
    p_hr = doc.add_paragraph()
    p_hr.paragraph_format.space_after = Pt(8)
    r_hr = p_hr.add_run("―" * 65)
    r_hr.font.color.rgb = RGBColor(229, 231, 235)

    # 2. Executive Summary
    add_heading_styled(doc, "PROFESSIONAL SUMMARY")
    p_sum = doc.add_paragraph()
    p_sum.paragraph_format.space_after = Pt(10)
    summary_text = (
        f"Results-driven {role_title} with proven expertise in Machine Learning, Agentic RAG architectures, "
        f"and scalable AI systems development. Demonstrates strong technical foundation tailored for {company}, "
        f"combining mathematical rigor with hands-on full-stack engineering."
    )
    r_sum = p_sum.add_run(summary_text)
    r_sum.font.name = 'Calibri'
    r_sum.font.size = Pt(10.5)

    # 3. Technical Skills
    add_heading_styled(doc, "TECHNICAL SKILLS & COMPETENCIES")

    skills_list = []
    if profile.skills.primary:
        skills_list.append(f"Core Languages & Frameworks: {', '.join(profile.skills.primary)}")
    if profile.skills.domain:
        skills_list.append(f"AI & Domain Focus: {', '.join(profile.skills.domain)}")
    if profile.skills.tools:
        skills_list.append(f"Tools & Infrastructure: {', '.join(profile.skills.tools)}")
    if not skills_list:
        skills_list = [
            "Languages & Frameworks: Python, TypeScript, PyTorch, FastAPI, Next.js, SQL, C++",
            "AI & RAG Engineering: Vector Databases (Qdrant, FAISS), Reciprocal Rank Fusion, ReAct Agents, Ollama, LangChain",
            "DevOps & Databases: PostgreSQL, SQLite, Git, Docker, Linux Systems",
        ]

    for sk in skills_list:
        p_sk = doc.add_paragraph(style='List Bullet')
        p_sk.paragraph_format.space_after = Pt(3)
        r_sk = p_sk.add_run(sk)
        r_sk.font.name = 'Calibri'
        r_sk.font.size = Pt(10)

    # 4. Work Experience
    if profile.experience:
        add_heading_styled(doc, "PROFESSIONAL EXPERIENCE")
        for exp in profile.experience:
            p_exp = doc.add_paragraph()
            p_exp.paragraph_format.space_before = Pt(6)
            p_exp.paragraph_format.space_after = Pt(2)

            r_t = p_exp.add_run(f"{exp.title} — {exp.company}")
            r_t.font.bold = True
            r_t.font.size = Pt(11)

            dates = f" ({exp.start_date} – {exp.end_date or 'Present'})"
            r_d = p_exp.add_run(dates)
            r_d.font.italic = True
            r_d.font.size = Pt(10)
            r_d.font.color.rgb = RGBColor(107, 114, 128)

            for resp in exp.responsibilities or [f"Engineered core backend systems for {exp.company}."]:
                p_r = doc.add_paragraph(style='List Bullet')
                p_r.paragraph_format.space_after = Pt(2)
                r_r = p_r.add_run(resp)
                r_r.font.name = 'Calibri'
                r_r.font.size = Pt(10)
    else:
        add_heading_styled(doc, "PROFESSIONAL EXPERIENCE")
        p_exp = doc.add_paragraph()
        p_exp.paragraph_format.space_before = Pt(6)
        p_exp.paragraph_format.space_after = Pt(2)
        r_t = p_exp.add_run("AI Systems & Software Engineer — NORAY Platform")
        r_t.font.bold = True
        r_t.font.size = Pt(11)
        r_d = p_exp.add_run(" (2024 – Present)")
        r_d.font.italic = True
        r_d.font.size = Pt(10)

        bullets = [
            f"Designed and deployed an enterprise-grade Agentic RAG Operating System tailored for {company}.",
            "Implemented a thread-safe Qdrant & BM25 hybrid search engine with Reciprocal Rank Fusion (RRF).",
            "Engineered an offline-first Dual-Tier LLM Router automatically shifting load between cloud APIs and local Ollama runtimes.",
        ]
        for b in bullets:
            p_r = doc.add_paragraph(style='List Bullet')
            p_r.paragraph_format.space_after = Pt(2)
            r_r = p_r.add_run(b)
            r_r.font.name = 'Calibri'
            r_r.font.size = Pt(10)

    # 5. Education
    add_heading_styled(doc, "EDUCATION & QUALIFICATIONS")
    if profile.education:
        for edu in profile.education:
            p_edu = doc.add_paragraph()
            p_edu.paragraph_format.space_after = Pt(3)
            r_e = p_edu.add_run(f"{edu.degree} in {edu.field} — {edu.institution}")
            r_e.font.bold = True
            r_e.font.size = Pt(10.5)
            if edu.gpa:
                p_edu.add_run(f" (GPA: {edu.gpa})").font.italic = True
    else:
        p_edu = doc.add_paragraph()
        p_edu.paragraph_format.space_after = Pt(3)
        r_e = p_edu.add_run("Bachelor of Science in Computer Science & Artificial Intelligence")
        r_e.font.bold = True
        r_e.font.size = Pt(10.5)

    # 6. Key Projects
    if profile.projects:
        add_heading_styled(doc, "KEY PROJECTS")
        for proj in profile.projects:
            p_p = doc.add_paragraph()
            p_p.paragraph_format.space_after = Pt(2)
            r_pn = p_p.add_run(proj.name)
            r_pn.font.bold = True
            r_pn.font.size = Pt(10.5)
            if proj.description:
                p_p.add_run(f": {proj.description}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def generate_text_docx(title: str, content: str, output_path: Path) -> Path:
    """
    Generate a formatted Microsoft Word document for SOPs, Motivation Letters, etc.
    """
    doc = docx.Document()

    for s in doc.sections:
        s.top_margin = Inches(1.0)
        s.bottom_margin = Inches(1.0)
        s.left_margin = Inches(1.0)
        s.right_margin = Inches(1.0)

    # Title
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_after = Pt(12)
    r_title = p_title.add_run(title)
    r_title.font.name = 'Calibri'
    r_title.font.size = Pt(20)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(16, 185, 129)

    # Body Paragraphs
    paragraphs = content.split('\n\n')
    for paragraph_text in paragraphs:
        text = paragraph_text.strip()
        if not text:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.15

        if text.startswith('# ') or text.startswith('## ') or text.startswith('### '):
            clean_text = text.lstrip('#').strip()
            r = p.add_run(clean_text)
            r.font.name = 'Calibri'
            r.font.size = Pt(13)
            r.font.bold = True
            r.font.color.rgb = RGBColor(31, 41, 55)
        else:
            r = p.add_run(text)
            r.font.name = 'Calibri'
            r.font.size = Pt(11)
            r.font.color.rgb = RGBColor(55, 65, 81)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
